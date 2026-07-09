import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


BASE_DIR = os.path.dirname(__file__)
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs")
DOC_PATH = os.path.join(OUTPUT_DIR, "Trabalho_NLP_Fase5_Monique_Silva_Medeiros.docx")


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    doc.add_paragraph(text)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)

    doc.add_paragraph("")
    return table


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CLASSIFICAÇÃO DE TEXTOS COM NLP\n")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Prova Substitutiva - Data Analytics Fase 5\n"
        "Dataset: BBC News (bbc-text.csv)"
    )

    doc.add_paragraph("")
    doc.add_paragraph("")

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run("Aluna: Monique Silva Medeiros\nRM: 362520")

    doc.add_page_break()

    add_heading(doc, "1. Introdução", level=1)
    add_paragraph(
        doc,
        "Este trabalho tem como objetivo implementar e comparar duas abordagens de "
        "Processamento de Linguagem Natural (NLP) para classificação automática de "
        "textos em categorias. A tarefa consiste em classificar notícias da BBC em "
        "cinco categorias: business, entertainment, politics, sport e tech."
    )
    add_paragraph(
        doc,
        "Foram exploradas duas soluções: (1) Deep Learning com embeddings Word2Vec "
        "pré-treinados e rede neural recorrente (LSTM); e (2) Word2Vec com Gensim, "
        "média dos vetores por documento e classificação com Regressão Logística. "
        "O objetivo final é comparar o desempenho das duas abordagens em termos de "
        "precisão, recall, F1-score e tempo de execução."
    )

    add_heading(doc, "2. Coleta de Dados", level=1)
    add_paragraph(
        doc,
        "Foi utilizado o dataset BBC News, disponibilizado no repositório da FIAP, "
        "contendo notícias rotuladas por categoria. O arquivo bbc-text.csv possui as "
        "colunas category (classe do documento) e text (conteúdo textual da notícia). "
        "O conjunto possui aproximadamente 2.225 documentos distribuídos entre as cinco categorias."
    )

    add_heading(doc, "3. Pré-processamento de Texto", level=1)
    add_paragraph(doc, "Antes do treinamento dos modelos, foi aplicado o seguinte pipeline:")
    for item in [
        "Conversão do texto para minúsculas;",
        "Tokenização com a biblioteca NLTK;",
        "Remoção de pontuação (mantendo apenas tokens alfabéticos);",
        "Remoção de stopwords em inglês.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "4. Metodologia", level=1)

    add_heading(doc, "4.1 Divisão dos Dados", level=2)
    add_paragraph(
        doc,
        "Os dados foram divididos em 80% para treino e 20% para teste, com divisão "
        "estratificada por categoria. O mesmo conjunto de teste foi utilizado para "
        "avaliar as duas abordagens."
    )

    add_heading(doc, "4.2 Abordagem 1 - Deep Learning (Word2Vec + LSTM)", level=2)
    add_paragraph(
        doc,
        "Foi treinado um modelo Word2Vec com Gensim no conjunto de treino. Os vetores "
        "obtidos foram utilizados como embeddings pré-treinados na camada inicial da "
        "rede neural. A arquitetura utilizada foi: Embedding (Word2Vec), LSTM (128 unidades), "
        "Dropout, Dense (64 neurônios), Dropout e camada de saída Softmax com 5 classes."
    )
    add_paragraph(
        doc,
        "O modelo foi treinado com otimizador Adam, função de perda sparse_categorical_crossentropy "
        "e Early Stopping para evitar overfitting. A avaliação foi feita com precision, recall, "
        "F1-score e matriz de confusão."
    )

    add_heading(doc, "4.3 Abordagem 2 - Word2Vec + Regressão Logística", level=2)
    add_paragraph(
        doc,
        "O mesmo Word2Vec treinado no conjunto de treino foi utilizado para representar cada "
        "documento pela média aritmética dos vetores das palavras. Em seguida, um classificador "
        "de Regressão Logística foi treinado para prever a categoria. As mesmas métricas da "
        "abordagem de Deep Learning foram calculadas no mesmo conjunto de teste."
    )

    add_heading(doc, "5. Resultados", level=1)

    add_heading(doc, "5.1 Comparativo Geral", level=2)
    add_table(
        doc,
        ["Modelo", "Acurácia", "Precision (macro)", "Recall (macro)", "F1-score (macro)", "Tempo"],
        [
            ["Deep Learning (Word2Vec + LSTM)", "60,9%", "0,587", "0,572", "0,534", "44,91 s"],
            ["Word2Vec + Regressão Logística", "92,4%", "0,924", "0,920", "0,922", "0,83 s"],
        ],
    )

    add_heading(doc, "5.2 Deep Learning - Métricas por Classe", level=2)
    add_table(
        doc,
        ["Categoria", "Precision", "Recall", "F1-score", "Support"],
        [
            ["business", "0,669", "0,833", "0,742", "102"],
            ["entertainment", "0,200", "0,039", "0,065", "77"],
            ["politics", "0,684", "0,643", "0,663", "84"],
            ["sport", "0,524", "0,971", "0,680", "102"],
            ["tech", "0,857", "0,375", "0,522", "80"],
        ],
    )

    add_heading(doc, "5.3 Word2Vec + Regressão Logística - Métricas por Classe", level=2)
    add_table(
        doc,
        ["Categoria", "Precision", "Recall", "F1-score", "Support"],
        [
            ["business", "0,912", "0,912", "0,912", "102"],
            ["entertainment", "0,957", "0,870", "0,912", "77"],
            ["politics", "0,897", "0,929", "0,912", "84"],
            ["sport", "0,944", "0,990", "0,967", "102"],
            ["tech", "0,911", "0,900", "0,906", "80"],
        ],
    )

    add_heading(doc, "5.4 Matrizes de Confusão", level=2)

    deep_cm = os.path.join(OUTPUT_DIR, "confusion_matrix_deep_learning.png")
    w2v_cm = os.path.join(OUTPUT_DIR, "confusion_matrix_word2vec.png")

    add_paragraph(doc, "Deep Learning (Word2Vec + LSTM):")
    if os.path.exists(deep_cm):
        doc.add_picture(deep_cm, width=Inches(5.5))
    else:
        add_paragraph(doc, "[Imagem não encontrada. Execute py main.py para gerar a matriz de confusão.]")

    doc.add_paragraph("")
    add_paragraph(doc, "Word2Vec + Regressão Logística:")
    if os.path.exists(w2v_cm):
        doc.add_picture(w2v_cm, width=Inches(5.5))
    else:
        add_paragraph(doc, "[Imagem não encontrada. Execute py main.py para gerar a matriz de confusão.]")

    add_heading(doc, "6. Análise dos Resultados", level=1)
    add_paragraph(
        doc,
        "A abordagem Word2Vec + Regressão Logística obteve desempenho significativamente "
        "superior, com acurácia de 92,4% e F1-score macro de 0,922, em 0,83 segundos. "
        "Já o modelo Deep Learning (Word2Vec + LSTM) alcançou acurácia de 60,9% e F1-score "
        "macro de 0,534, em 44,91 segundos."
    )
    add_paragraph(
        doc,
        "O modelo de Deep Learning apresentou bom desempenho nas categorias politics e tech, "
        "mas falhou completamente na classificação de entertainment e teve recall muito baixo "
        "em business. Isso indica que, para este dataset e configuração, a abordagem clássica "
        "com média de vetores Word2Vec e classificador linear foi mais eficaz."
    )

    add_heading(doc, "7. Conclusão", level=1)
    add_paragraph(
        doc,
        "Neste experimento com o dataset BBC News, a abordagem Word2Vec + Regressão Logística "
        "apresentou o melhor desempenho na tarefa de classificação de textos."
    )

    add_heading(doc, "Vantagens do Deep Learning (Word2Vec + LSTM)", level=2)
    for item in [
        "Capacidade de capturar dependências sequenciais e contexto entre palavras;",
        "Possibilidade de fine-tuning dos embeddings durante o treinamento;",
        "Potencial para escalar melhor com mais dados e textos mais complexos.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Desvantagens do Deep Learning", level=2)
    for item in [
        "Maior custo computacional e tempo de treinamento;",
        "Exige mais ajuste de hiperparâmetros;",
        "Menor interpretabilidade dos resultados;",
        "Desempenho inferior neste experimento específico.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Vantagens do Word2Vec + Regressão Logística", level=2)
    for item in [
        "Implementação mais simples e treinamento muito mais rápido;",
        "Excelente desempenho em classificação de notícias com vocabulário estável;",
        "Maior interpretabilidade do classificador linear;",
        "Resultados consistentes em todas as categorias.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Desvantagens do Word2Vec + Regressão Logística", level=2)
    for item in [
        "A média dos vetores perde informação sobre a ordem das palavras;",
        "Depende da qualidade do Word2Vec treinado no corpus;",
        "Pode ser limitado em tarefas que exigem compreensão profunda de contexto.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_paragraph(
        doc,
        "Conclusão final: para o dataset BBC News utilizado neste trabalho, a abordagem "
        "Word2Vec + Regressão Logística demonstrou ser a mais adequada, combinando alta "
        "precisão, bom equilíbrio entre precision e recall, e eficiência computacional."
    )

    add_heading(doc, "8. Referências", level=1)
    for item in [
        "FIAP. Repositório Pos Tech DTAT - Dataset BBC News.",
        "Bird, S.; Loper, E.; Klein, E. NLTK: Natural Language Toolkit.",
        "Řehůřek, R.; Sojka, P. Gensim - Statistical Semantics in Python.",
        "Mikolov, T. et al. Efficient Estimation of Word Representations in Vector Space (Word2Vec).",
        "Chollet, F. Keras / TensorFlow - Deep Learning framework.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(DOC_PATH)
    print(f"Documento gerado com sucesso: {DOC_PATH}")


if __name__ == "__main__":
    main()
